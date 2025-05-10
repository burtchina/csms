#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
将Markdown文档转换为Word文档
"""

import os
import re
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_file, docx_file):
    """将Markdown文件转换为Word文档"""
    
    print(f"正在转换 {md_file} 到 {docx_file}")
    
    # 读取Markdown文件内容
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档标题
    title_match = re.search(r'^# (.*?)$', md_content, re.MULTILINE)
    if title_match:
        title = title_match.group(1)
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 处理Markdown内容
    # 移除代码块和Mermaid图表标记，因为无法直接转换
    md_content = re.sub(r'```mermaid[\s\S]*?```', '[Mermaid图表]', md_content)
    md_content = re.sub(r'```[\s\S]*?```', '[代码块]', md_content)
    
    # 分段处理
    sections = md_content.split('\n## ')
    
    # 处理第一部分（不包含标题的内容）
    if sections:
        first_section = sections[0]
        # 移除主标题行
        first_section = re.sub(r'^# .*?\n', '', first_section)
        
        # 添加内容
        for line in first_section.split('\n'):
            if line.strip():
                # 检查是否是标题
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=3)
                else:
                    doc.add_paragraph(line)
    
    # 处理其他部分
    for section in sections[1:]:
        section_parts = section.split('\n', 1)
        if len(section_parts) > 0:
            section_title = section_parts[0]
            doc.add_heading(section_title, level=1)
            
            if len(section_parts) > 1:
                section_content = section_parts[1]
                
                for line in section_content.split('\n'):
                    if line.strip():
                        # 检查是否是标题
                        if line.startswith('### '):
                            doc.add_heading(line[4:], level=2)
                        elif line.startswith('#### '):
                            doc.add_heading(line[5:], level=3)
                        else:
                            doc.add_paragraph(line)
    
    # 保存文档
    doc.save(docx_file)
    print(f"文档已保存到 {docx_file}")

if __name__ == "__main__":
    md_file = "doc/system_architecture_analysis.md"
    docx_file = "doc/system_architecture_analysis.docx"
    
    md_to_docx(md_file, docx_file) 